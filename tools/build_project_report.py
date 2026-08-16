from __future__ import annotations

"""Create the internship project report as a DOCX document.

The document intentionally uses the reference report's technical-report rhythm:
numbered chapters, compact tables, figure captions, explicit experiment notes,
and a separate limitations/conclusion section.  Project facts are written from
the repository README, source code, test results, and the recorded experiment
summaries; missing local model/audio artefacts are never silently treated as
new measurements.
"""

from pathlib import Path
import math

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "outputs" / "HAVALIMANI_GURULTU_TESPIT_PROJE_RAPORU.docx"
WORK = ROOT / ".report_work"
FIGURES = WORK / "generated_figures"

NAVY = "1F4E79"
BLUE = "2E75B6"
PALE_BLUE = "D9EAF7"
LIGHT_BLUE = "EEF5FB"
GRAY = "595959"
LIGHT_GRAY = "F2F2F2"
MID_GRAY = "D9E1F2"
GREEN = "70AD47"
ORANGE = "ED7D31"
RED = "C00000"
BLACK = "202020"


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def font_path() -> str | None:
    for candidate in (Path("C:/Windows/Fonts/arial.ttf"), Path("C:/Windows/Fonts/calibri.ttf")):
        if candidate.exists():
            return str(candidate)
    return None


def load_font(size: int, bold: bool = False):
    candidates = [
        Path("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size)
    return ImageFont.load_default()


def hex_fill(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_border(cell, **kwargs) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge not in kwargs:
            continue
        tag = "w:%s" % edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        for key in ("val", "sz", "space", "color"):
            if key in kwargs[edge]:
                element.set(qn("w:%s" % key), str(kwargs[edge][key]))


def set_cell_margins(cell, top=90, start=100, bottom=90, end=100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn("w:%s" % name))
        if node is None:
            node = OxmlElement("w:%s" % name)
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[float]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            if index < len(widths):
                cell.width = Inches(widths[index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            set_cell_border(
                cell,
                top={"val": "single", "sz": 4, "color": "D9E2F3"},
                bottom={"val": "single", "sz": 4, "color": "D9E2F3"},
                left={"val": "single", "sz": 4, "color": "D9E2F3"},
                right={"val": "single", "sz": 4, "color": "D9E2F3"},
            )


def repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def keep_with_next(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    keep = OxmlElement("w:keepNext")
    p_pr.append(keep)


def set_run_font(run, name="Arial", size=10.5, color=BLACK, bold=False, italic=False):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    run.bold = bold
    run.italic = italic


def configure_document(doc: Document) -> None:
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(0.98)
    sec.bottom_margin = Inches(0.88)
    sec.left_margin = Inches(0.98)
    sec.right_margin = Inches(0.98)
    sec.header_distance = Inches(0.45)
    sec.footer_distance = Inches(0.45)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = rgb(BLACK)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.08

    for style_name, size, color, before, after in (
        ("Heading 1", 16, NAVY, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 11.5, GRAY, 9, 4),
    ):
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Caption" not in styles:
        caption = styles.add_style("Caption", WD_STYLE_TYPE.PARAGRAPH)
    else:
        caption = styles["Caption"]
    caption.font.name = "Arial"
    caption._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Arial")
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = rgb(GRAY)
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(8)


def add_header_footer(doc: Document) -> None:
    sec = doc.sections[0]
    header = sec.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("Havalimanı Gürültü Tespit ve Ses Tanıma Sistemi")
    set_run_font(run, size=8.5, color=GRAY)
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Staj Projesi Raporu  •  ")
    set_run_font(run, size=8.5, color=GRAY)
    run = footer.add_run()
    set_run_font(run, size=8.5, color=GRAY)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)


def add_title(doc, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    keep_with_next(p)
    return p


def add_body(doc, text: str, *, bold_lead: str | None = None, italic=False):
    p = doc.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.25)
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        set_run_font(r, bold=True)
        r = p.add_run(text[len(bold_lead):])
        set_run_font(r, italic=italic)
    else:
        r = p.add_run(text)
        set_run_font(r, italic=italic)
    return p


def add_small(doc, text: str, *, color=GRAY, italic=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    set_run_font(r, size=9, color=color, italic=italic)
    return p


def add_bullets(doc, items: list[str]):
    for item in items:
        p = doc.add_paragraph(style="Normal")
        p.style = doc.styles["Normal"]
        p.paragraph_format.left_indent = Inches(0.28)
        p.paragraph_format.first_line_indent = Inches(-0.16)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run("• ")
        set_run_font(r, bold=True, color=BLUE)
        r = p.add_run(item)
        set_run_font(r)


def add_table(doc, headers: list[str], rows: list[list[str]], widths: list[float], *, font_size=8.7):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0]
    repeat_header(hdr)
    for index, text in enumerate(headers):
        cell = hdr.cells[index]
        hex_fill(cell, NAVY)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(str(text))
        set_run_font(r, size=font_size, color="FFFFFF", bold=True)
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cell = cells[index]
            hex_fill(cell, LIGHT_BLUE if row_index % 2 == 0 else "FFFFFF")
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if index == 0 else WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(value))
            set_run_font(r, size=font_size, color=BLACK)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_note(doc, text: str, *, color=BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    hex_fill(cell, "F5F9FD")
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    set_run_font(r, size=9.2, color=color, italic=True)
    set_table_geometry(table, [6.45])
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_caption(doc, text: str):
    p = doc.add_paragraph(style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run_font(r, size=9, color=GRAY, italic=True)
    return p


def add_figure(doc, path: Path, caption: str, *, width=6.1):
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    add_caption(doc, caption)


def save_bar_chart(path: Path, title: str, labels: list[str], values: list[float], *, max_value=100, color=BLUE, suffix="%"):
    path.parent.mkdir(parents=True, exist_ok=True)
    width, height = 1320, 650
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    title_font = load_font(28, True)
    label_font = load_font(21)
    value_font = load_font(22, True)
    draw.text((50, 24), title, fill="#1F4E79", font=title_font)
    left, right, top, bottom = 115, 1250, 105, 560
    draw.line((left, top, left, bottom), fill="#9EADBA", width=2)
    draw.line((left, bottom, right, bottom), fill="#9EADBA", width=2)
    for tick in range(0, 101, 20):
        y = bottom - (bottom - top) * tick / max_value
        draw.line((left, y, right, y), fill="#E7EDF2", width=1)
        draw.text((45, y - 12), f"{tick}{suffix}", fill="#59636B", font=label_font, anchor="rm")
    slot = (right - left) / max(1, len(values))
    bar_width = slot * 0.62
    for index, value in enumerate(values):
        x0 = left + slot * index + (slot - bar_width) / 2
        x1 = x0 + bar_width
        y = bottom - (bottom - top) * value / max_value
        draw.rounded_rectangle((x0, y, x1, bottom), radius=8, fill="#" + color.lstrip("#"))
        draw.text(((x0 + x1) / 2, y - 25), f"{value:.1f}{suffix}", fill="#202020", font=value_font, anchor="ms")
        words = labels[index].split(" ")
        if len(words) > 2:
            midpoint = max(1, len(words) // 2)
            lines = [" ".join(words[:midpoint]), " ".join(words[midpoint:])]
        else:
            lines = [labels[index]]
        for line_index, line in enumerate(lines):
            draw.text(((x0 + x1) / 2, bottom + 18 + line_index * 25), line, fill="#59636B", font=label_font, anchor="ma")
    img.save(path)


def save_grouped_chart(path: Path, title: str, labels: list[str], series: list[tuple[str, list[float], str]]):
    path.parent.mkdir(parents=True, exist_ok=True)
    width, height = 1420, 720
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    title_font = load_font(28, True)
    label_font = load_font(20)
    value_font = load_font(18, True)
    draw.text((50, 24), title, fill="#1F4E79", font=title_font)
    left, right, top, bottom = 120, 1320, 115, 600
    draw.line((left, top, left, bottom), fill="#9EADBA", width=2)
    draw.line((left, bottom, right, bottom), fill="#9EADBA", width=2)
    for tick in range(0, 101, 20):
        y = bottom - (bottom - top) * tick / 100
        draw.line((left, y, right, y), fill="#E7EDF2", width=1)
        draw.text((50, y - 12), f"{tick}%", fill="#59636B", font=label_font, anchor="rm")
    slot = (right - left) / max(1, len(labels))
    group_width = slot * 0.72
    bar_width = group_width / max(1, len(series)) - 8
    for label_index, label in enumerate(labels):
        base = left + slot * label_index + (slot - group_width) / 2
        for series_index, (_, values, color) in enumerate(series):
            value = values[label_index]
            x0 = base + series_index * (bar_width + 8)
            x1 = x0 + bar_width
            y = bottom - (bottom - top) * value / 100
            draw.rectangle((x0, y, x1, bottom), fill="#" + color.lstrip("#"))
            draw.text(((x0 + x1) / 2, y - 16), f"{value:.0f}", fill="#202020", font=value_font, anchor="ms")
        draw.text((left + slot * label_index + slot / 2, bottom + 18), label, fill="#59636B", font=label_font, anchor="ma")
    legend_x = left
    for name, _, color in series:
        draw.rectangle((legend_x, 655, legend_x + 22, 677), fill="#" + color.lstrip("#"))
        draw.text((legend_x + 30, 653), name, fill="#59636B", font=label_font)
        legend_x += 210
    img.save(path)


def make_figures() -> dict[str, Path]:
    FIGURES.mkdir(parents=True, exist_ok=True)
    figures = {
        "main_models": FIGURES / "main_model_comparison.png",
        "subtypes": FIGURES / "subtype_metrics.png",
        "shazam": FIGURES / "shazam_pipeline.png",
        "tests": FIGURES / "test_evidence.png",
        "window": FIGURES / "window_voting.png",
    }
    save_grouped_chart(
        figures["main_models"],
        "Altı üst kategori bağımsız testinde model karşılaştırması",
        ["Masking", "Contrastive", "Masking +\nContrastive", "Hibrit +\nSampler", "Hibrit +\nOTHER uzmanı"],
        [("Doğruluk", [71.60, 72.10, 72.27, 73.28, 73.78], BLUE), ("Macro-F1", [68.34, 68.82, 69.35, 70.66, 71.00], ORANGE)],
    )
    save_grouped_chart(
        figures["subtypes"],
        "Alt tür deneylerinin test Macro-F1 sonuçları",
        ["TRAFFIC", "OTHER", "Uçak tipi"],
        [("Test Macro-F1", [98.14, 79.04, 50.20], GREEN)],
    )
    save_bar_chart(
        figures["shazam"],
        "Shazam parmak izi katmanı: kayıt ve hash ölçeği",
        ["Uçak referansı", "Birleşik katalog", "Kategori hash"],
        [851437 / 26000, 2585412 / 26000, 1733975 / 26000],
        max_value=110,
        color=BLUE,
        suffix="k",
    )
    save_grouped_chart(
        figures["tests"],
        "Çoklu pencere oylaması doğrulama özeti",
        ["Hedefli", "Tam proje"],
        [("Başarılı", [13, 88], GREEN), ("Başarısız", [0, 0], RED)],
    )
    save_bar_chart(
        figures["window"],
        "30,55 saniyelik örnekte seçilen pencere başlangıçları",
        ["P1 0 s", "P2 7,5 s", "P3 15 s", "P4 20 s", "P5 25,55 s"],
        [1, 1, 1, 1, 1],
        max_value=1,
        color=ORANGE,
        suffix="",
    )
    return figures


def add_page_break(doc: Document):
    # A standalone break paragraph can be pushed to the next page when the
    # previous table already fills the page, producing an unwanted blank page.
    # page_break_before lets Word start the next real paragraph on a new page.
    p = doc.add_paragraph()
    p.paragraph_format.page_break_before = True


def build_report() -> Path:
    figures = make_figures()
    doc = Document()
    configure_document(doc)
    add_header_footer(doc)

    # Kapak
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("STAJ PROJESİ RAPORU")
    set_run_font(r, size=15, color=BLUE, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Havalimanı Gürültü Tespit ve Ses Tanıma Sistemi")
    set_run_font(r, size=24, color=NAVY, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Üst kategori sınıflandırması, uçak parmak izi eşleştirmesi ve çoklu pencere oylaması")
    set_run_font(r, size=12, color=GRAY, italic=True)
    for _ in range(3):
        doc.add_paragraph()
    meta = [
        ["Hazırlayan", "[Ad Soyad]"],
        ["Bölüm / Program", "[Bölüm bilgisi]"],
        ["Kurum", "[Staj yapılan kurum]"],
        ["Rapor tarihi", "16 Ağustos 2026"],
        ["Proje klasörü", "C:\\Airport_Noise_Detection-main"],
    ]
    add_table(doc, ["Bilgi", "Açıklama"], meta, [1.75, 4.7], font_size=9.5)
    add_note(doc, "Bu rapor, gönderilen örnek raporun teknik anlatım düzeni ve ayrıntı seviyesi örnek alınarak hazırlanmıştır. Örnekteki kurumsal/ödev talimatları proje gerçeği olarak kopyalanmamış; teknik içerik bu projeye ait kaynak kod, README, test sonuçları ve deney özetlerinden oluşturulmuştur.")
    add_page_break(doc)

    # Özet ve içerik özeti
    add_title(doc, "ÖZET", 1)
    add_body(doc, "Bu staj projesinde havalimanı çevresindeki sesleri önce altı ana kategoriye ayıran, ardından uygun kategori için alt tür veya kayıt eşleştirmesi yapan bir masaüstü sistemi geliştirdim. Projenin ilk bakışta sınıflandırma problemi gibi görünmesine rağmen uygulamada iki ayrı karar türü olduğunu gördüm: Bir sesin genel olarak uçak, trafik veya konuşma olup olmadığını öğrenilmiş model tahminiyle bulmak; katalogda daha önce insan tarafından doğrulanmış bir uçak kaydı varsa bunu akustik parmak iziyle kesin eşleştirmek. Bu ayrım raporun tamamında özellikle korunmuştur.")
    add_body(doc, "Çalışmanın en somut çıktıları; dosyadan ses analizi yapan PyQt tabanlı arayüz, EfficientNet/CNN/SVM ve isteğe bağlı BEATs kanallarını bir araya getiren sınıflandırma orkestrasyonu, uçak güvenlik kapısı, Shazam tarzı fingerprint kataloğu, insan onaylı referans laboratuvarı, salt-okunur SQLite kanıt ekranı ve bu raporun konusu olan çoklu pencere oylamasıdır. Çoklu pencere değişikliğiyle uzun bir ses kaydının yalnızca başından karar vermek yerine kaydın başını, ortasını ve sonunu temsil eden en fazla beş tam pencere işlenmiş ve pencere tahminleri çoğunluk oyuyla birleştirilmiştir.")
    add_table(doc, ["Sonuç alanı", "Ölçülen sonuç", "Kapsam / yorum"], [
        ["Ana kategori", "%95,83", "648 kayıtlık uçtan uca akışta 621 doğru"],
        ["Alt tür v2", "%86,57", "Aynı 648 kayıtta 561 doğru uçtan uca sonuç"],
        ["TRAFFIC alt türü", "%98,14 Macro-F1", "1.464 / 312 / 324 train-val-test"],
        ["OTHER alt türü", "%79,04 Macro-F1", "Aynı split düzeninde altı sınıf"],
        ["Uçak tipi modeli", "%50,20 Macro-F1", "432 test kaydı, sekiz sınıf; sınırlı genelleme"],
        ["Çoklu pencere testleri", "13/13 ve 88/88", "Hedefli ve tam test kümeleri başarılı"],
    ], [1.55, 1.55, 3.35])
    add_note(doc, "Yüzdeler aynı veri kümesini veya aynı problemi temsil etmiyor. Rapor boyunca her metrik yanında test kapsamını ayrıca belirttim; özellikle uçtan uca sonuç ile tek bir modelin Macro-F1 değeri birbirine eşitlenmemelidir.")
    add_page_break(doc)

    add_title(doc, "İÇİNDEKİLER VE RAPORU OKUMA NOTU", 1)
    add_table(doc, ["Bölüm", "İçerik"], [
        ["1", "Projenin amacı, problem ve başarı ölçütleri"],
        ["2", "Sistem mimarisi ve karar akışı"],
        ["3", "Veri kaynakları, etiketleme ve veri sızıntısı kontrolü"],
        ["4", "Ön işleme ve model kanalları"],
        ["5", "Shazam tarzı akustik parmak izi"],
        ["6", "Model deneyleri ve karşılaştırmalar"],
        ["7", "Başarısız denemeler ve bunlardan çıkarılan dersler"],
        ["8", "Çoklu pencere oylamasının uygulanması ve test edilmesi"],
        ["9", "Uçtan uca uygulama davranışı ve kullanıcı arayüzü"],
        ["10", "Yeniden üretilebilirlik, sınırlılıklar ve riskler"],
        ["11", "Sonuç ve gelecek çalışma planı"],
        ["12", "Terimler, dosyalar ve kaynak notları"],
    ], [0.8, 5.65])
    add_body(doc, "Raporun amacı yalnızca iyi çalışan sonuçları sıralamak değildir. Bir sonraki kişinin projeyi açtığında hangi dosyanın ne yaptığını, hangi deneyin neden denendiğini, hangi sonucun üretime alınmadığını ve hangi sonuçların daha fazla veriye ihtiyaç duyduğunu anlayabilmesidir. Bu nedenle başarısız deneyler ve ölçüm kapsamları özellikle saklanmıştır.")
    add_page_break(doc)

    # 1
    add_title(doc, "1. PROJENİN AMACI, PROBLEM TANIMI VE KAPSAMI", 1)
    add_title(doc, "1.1. Problem tanımı", 2)
    add_body(doc, "Havalimanı çevresindeki ses kayıtları tek bir olaydan oluşmayabilir. Bir kaydın ilk saniyelerinde rüzgâr, ortasında araç trafiği, son kısmında ise uçak motoru duyulabilir. Mikrofon konumu, kayıt seviyesi, arka plan gürültüsü ve sesin kayda göre uzaklığı da değişir. Bu yüzden yalnızca tek bir kısa parçadan tahmin almak, özellikle uzun dosyalarda kararsız veya aşırı kendinden emin sonuçlar üretebilir.")
    add_body(doc, "İkinci problem, uçak tipini genel ses benzerliğinden tahmin etmek ile belirli bir referans kaydın tekrarını bulmanın aynı şey olmamasıdır. Öğrenilmiş model yeni bir fiziksel uçağa genelleme yapmaya çalışır; fingerprint katmanı ise katalogda bulunan doğrulanmış kaydı arar. Projede bu iki yaklaşımın sonucunu aynı etiket gibi sunmamak için yöntem bilgisi kullanıcıya gösterilir.")
    add_title(doc, "1.2. Projenin hedefleri", 2)
    add_bullets(doc, [
        "WAV veya mikrofon girdisini altı ana sınıftan birine ayırmak: AIRCRAFT, AMBIENT, OTHER, SPEECH, TRAFFIC ve WIND.",
        "AIRCRAFT sonucunu bir güvenlik kapısından geçirmek ve uygun olduğunda önce Shazam tarzı kayıt eşleştirmesi yapmak.",
        "TRAFFIC ve OTHER için mevcut alt tür modellerini, uçak için ise hem fingerprint hem öğrenilmiş model yolunu kullanmak.",
        "Uzun kayıtlarda tek pencere yanlılığını azaltmak için en fazla beş temsilci pencereyi çoğunluk oyu ile birleştirmek.",
        "Model tahminini, kullanılan yöntemi, pencere oylarını ve insan onayını izlenebilir biçimde saklamak.",
        "Başarılı olduğu kadar başarısız deneyleri de raporlayarak hangi iddianın hangi veri kapsamıyla sınırlı olduğunu göstermek.",
    ])
    add_title(doc, "1.3. Başarı ölçütleri", 2)
    add_table(doc, ["Ölçüt", "Hedeflenen davranış", "Kontrol yöntemi"], [
        ["Ana kategori", "Gürültülü kayıtta altı sınıftan makul bir seçim", "Bağımsız test ve uçtan uca 648 kayıt"],
        ["Uzun kayıt", "Baş/orta/son bölümlerinin karara katılması", "Pencere başlangıçları ve vote_counts"],
        ["Uçak güvenliği", "Uçak olmayan kayıtlarda alt tür zorlamamak", "False positive takibi ve guard testleri"],
        ["Fingerprint", "Katalogdaki doğrulanmış kaydı açıklanabilir biçimde bulmak", "Hash hizalanması ve skor eşiği"],
        ["İzlenebilirlik", "Sonuçla birlikte yöntem ve güveni göstermek", "GUI metadatası, JSON/SQLite kanıtı"],
        ["Sürdürülebilirlik", "Kurulum ve test adımlarının başkası tarafından tekrarlanabilmesi", "README, testler ve model yokluğu için kontrollü kapanış"],
    ], [1.25, 2.95, 2.25])
    add_page_break(doc)

    # 2
    add_title(doc, "2. SİSTEM MİMARİSİ VE KARAR AKIŞI", 1)
    add_title(doc, "2.1. Genel akış", 2)
    add_body(doc, "Sistem kaynak kodda tek bir model fonksiyonundan ibaret değildir. Ses yükleme, pencere seçimi, özellik çıkarımı, model tahmini, pencere oylaması, uçak güvenlik kontrolü, fingerprint veya alt tür yolu ve GUI sunumu ayrı sorumluluklar olarak ele alınmıştır. Böylece bir katmanda yapılan değişikliğin diğer katmanın mantığını sessizce bozması zorlaştırılmıştır.")
    add_table(doc, ["Sıra", "Katman", "Görevi", "Çıktısı"], [
        ["1", "Girdi", "WAV/mikrofonu ve örnekleme oranını alır", "Yüklü dalga biçimi"],
        ["2", "Pencereleme", "5 saniyelik pencereleri 2,5 saniye adımla seçer; en fazla 5", "Pencere dizisi ve başlangıçları"],
        ["3", "Ana model", "EfficientNet, CNN, SVM veya isteğe bağlı BEATs ile kategori tahmini", "Pencere başına sınıf/olasılık"],
        ["4", "Oylama", "Çoğunluk oyunu, eşitlikte ortalama olasılıkla çözer", "Kazanan, oy sayıları, güven"],
        ["5", "Uçak kapısı", "AIRCRAFT dışındaki sonucu uçak alt türüne taşımayı engeller", "İzin/verilmez kararı"],
        ["6", "Alt sonuç", "Uçakta fingerprint; diğerlerinde uygun alt tür modeli", "Alt tür veya açıklanmış fallback"],
        ["7", "Kullanıcı katmanı", "Sonucu ve yöntemini GUI'de gösterir", "İnsan tarafından okunabilir rapor"],
    ], [0.55, 1.25, 3.1, 1.55])
    add_title(doc, "2.2. Dosya ve modül sorumlulukları", 2)
    add_table(doc, ["Dosya", "Sorumluluk", "Bu rapordaki karşılığı"], [
        ["noise_detector.py", "Ana kategori, model orkestrasyonu, GUI analiz akışı ve fingerprint entegrasyonu", "Bölüm 2, 5, 8, 9"],
        ["window_voting.py", "Pencere seçimi ve olasılık oylaması", "Bölüm 8"],
        ["gui_main.py", "PyQt6 arayüzü ve kullanıcı aksiyonları", "Bölüm 9"],
        ["aircraft_fingerprint.py", "STFT tepeleri, hash üretimi, katalog sorgusu", "Bölüm 5"],
        ["test_window_voting.py", "Pencereleme ve bütünleşme testleri", "Bölüm 8"],
        ["train_*.py / dataset_builder.py", "Veri hazırlama ve model eğitim araçları", "Bölüm 3, 4, 6"],
        ["README.md", "Kurulum, veri/model artefaktı ve ana deney özeti", "Bölüm 10, 12"],
    ], [1.75, 3.25, 1.45])
    add_note(doc, "Kodda model ağırlıkları, ham sesler ve SQLite indeksleri Git'e alınmamıştır. Rapor bu nedenle hem kaynak kodu hem de daha önce kaydedilmiş deney özetlerini birlikte okur; temiz klonun tek başına tüm metrikleri yeniden üretmesi beklenmemelidir.")
    add_page_break(doc)

    # 3
    add_title(doc, "3. VERİ KAYNAKLARI, ETİKETLEME VE VERİ SIZINTISI", 1)
    add_title(doc, "3.1. Kategori taksonomisi", 2)
    add_table(doc, ["Etiket", "Anlamı", "Alt yol"], [
        ["AIRCRAFT", "Uçak, motor, kalkış ve iniş sesleri", "Fingerprint veya uçak tipi modeli"],
        ["AMBIENT", "Genel çevre ve havalimanı ortamı", "Zorunlu alt tür yok"],
        ["OTHER", "Diğer çevresel/hayvansal sesler", "Kedi, karga, köpek, papağan, tavus kuşu, serçe"],
        ["SPEECH", "Konuşma ve anons", "Zorunlu alt tür yok"],
        ["TRAFFIC", "Kara ulaşımı ve trafik", "Bisiklet, otobüs, otomobil, motosiklet, tren, kamyon"],
        ["WIND", "Rüzgâr ve mikrofondaki hava akışı", "Zorunlu alt tür yok"],
    ], [1.15, 3.15, 2.15])
    add_title(doc, "3.2. Veri bölme yaklaşımı", 2)
    add_body(doc, "Aynı kaydın farklı kırpımlarını train ve test tarafına dağıtmak ölçümü olduğundan iyi gösterebilir. Bu nedenle özellikle uçak tipi çalışmasında fiziksel uçak veya kaynak kayıt gruplarının ayrılması gerektiği kabul edilmiştir. Proje yönergesinde bir sınıfın üretim sınıfı sayılması için en az 12, tercihen 30–50 farklı ana kayda ulaşması gerektiği belirtilmiştir; bu ölçüt, az örnekli uçak tiplerinin neden temkinli raporlandığını açıklamaktadır.")
    add_table(doc, ["Deney", "Train", "Validation", "Test / bağımsız", "Not"], [
        ["Ana 6 sınıf", "2.271", "400", "595", "Bağımsız test; final uzman deneyinin özeti"],
        ["Kategori uçtan uca", "—", "—", "648", "Arayüz/akış ölçümü; aynı şey model F1'i değildir"],
        ["TRAFFIC alt türü", "1.464", "312", "324", "6 sınıf"],
        ["OTHER alt türü", "1.464", "312", "324", "6 sınıf"],
        ["Uçak tipi BEATs", "1.952", "416", "432", "8 sınıf"],
        ["41 sınıf uçak denemesi", "Karma", "Karma", "Fiziksel uçak holdout", "Genelleme başarısız"],
    ], [1.5, 0.8, 0.9, 1.15, 2.1])
    add_title(doc, "3.3. Veri sızıntısı açısından alınan kararlar", 2)
    add_bullets(doc, [
        "Test kayıtları aktif Shazam indeksine eklenmemiştir; aksi durumda eşleştirme sonucu öğrenilmiş bir genelleme gibi görünürdü.",
        "Uçak referansında fiziksel uçak, kaynak kayıt ve split ilişkisi ayrı izlenmeye çalışılmıştır.",
        "Model seçimi bağımsız test sonucuna bakılarak değil, validation Macro-F1 üzerinden yapılmıştır.",
        "Fingerprint katalog verisi ile öğrenilmiş sınıflandırma verisi kavramsal olarak ayrılmıştır.",
        "Az sayıda fiziksel uçak içeren sınıflarda yüksek validation sonucunun genelleme anlamına gelmediği ayrıca sınanmıştır.",
    ])
    add_page_break(doc)

    # 4
    add_title(doc, "4. ÖN İŞLEME VE MODEL KANALLARI", 1)
    add_title(doc, "4.1. Ortak ses ön işleme", 2)
    add_body(doc, "Ses kayıtları model kanalına girmeden önce tek kanala indirgenir, örnekleme oranı modelin beklediği değere getirilir ve sabit uzunluklu parçalara ayrılır. Çoklu pencere değişikliğinin temel noktası, bu sabit parçanın artık kaydın yalnızca başından alınmamasıdır. Kısa kayıtlar sıfırla doldurularak tek pencere olarak işlenir; uzun kayıtlar ise başlangıç, orta ve son bölümleri temsil edecek biçimde en fazla beş pencereye düşürülür.")
    add_title(doc, "4.2. SVM özellikleri", 2)
    add_body(doc, "SVM hattında zaman alanı ve frekans alanı istatistiklerinden oluşan 264 boyutlu bir özellik vektörü kullanılmıştır. Bu yaklaşımın avantajı hafif olması ve CPU üzerinde hızlı çalışmasıdır. Dezavantajı ise ham spektrogramdaki ayrıntılı zaman-frekans desenlerini CNN veya EfficientNet kadar doğrudan kullanamamasıdır.")
    add_table(doc, ["Kanal", "Girdi", "Güçlü yönü", "Sınırı"], [
        ["SVM", "264 boyutlu el yapımı özellik", "Hafif, hızlı, açıklanabilir bir baseline", "Karmaşık ses örüntülerinde sınırlı"],
        ["AirportCNN", "Mel/spektrogram görüntüsü", "Yerel zaman-frekans desenlerini öğrenir", "Dış testte sınıf davranışı kararsız kaldı"],
        ["EfficientNet-B0", "RGB'ye dönüştürülmüş Mel görüntüsü", "Transfer öğrenme ve daha güçlü görsel temsil", "Model ağırlığı ve GPU ihtiyacı"],
        ["BEATs", "Önceden eğitilmiş ses temsili + başlık", "Genel ses bilgisinden yararlanır", "Uçak tipi sınıfında genelleme sınırlı"],
        ["AST / diğer pilotlar", "Transformer tabanlı ses temsili", "Fiziksel uçak gruplamasını sınamak için yararlı", "Üretim yoluna alınmış tek çözüm değil"],
    ], [1.15, 2.0, 2.05, 1.25])
    add_title(doc, "4.3. Mel spektrogram ve RGB dönüşümü", 2)
    add_body(doc, "CNN ve EfficientNet kanallarında ses, frekans eksenini Mel ölçeğine taşıyan bir spektrograma çevrilmiştir. EfficientNet görüntü tabanlı bir omurga olduğu için tek kanallı ses temsili üç kanala kopyalanarak RGB biçiminde kullanılmıştır. Buradaki RGB ifadesi renkli bir fotoğraf anlamına gelmez; aynı akustik bilginin görüntü omurgasına uygun tensör biçiminde taşınmasıdır.")
    add_note(doc, "Ön işleme ayarları deneyler arasında değişebildiği için raporda her sonuç için veri kapsamı öne çıkarılmıştır. Bir modelin daha iyi çıkması yalnızca mimarisinden değil; crop, sampler, class weight, focal loss ve fiziksel grup ayrımından da etkilenebilir.")
    add_page_break(doc)

    # 5
    add_title(doc, "5. SHAZAM TARZI AKUSTİK PARMAK İZİ", 1)
    add_title(doc, "5.1. Neden parmak izi katmanı gerekliydi?", 2)
    add_body(doc, "Uçak tipini yalnızca sınıflandırıcıya bıraktığımda model, eğitimde gördüğü seslerin istatistiksel benzerliğinden bir öneri üretmektedir. Katalogda doğrulanmış bir uçak kaydının daha kısa veya gürültülü bir parçası varsa, bu parçayı doğrudan kayıtla eşleştirmek daha açıklanabilir bir sonuç verir. Bu yüzden fingerprint katmanı “uçak tipini tahmin eden ikinci bir sınıflandırıcı” olarak değil, doğrulanmış referans araması olarak tasarlanmıştır.")
    add_title(doc, "5.2. Hash üretim adımları", 2)
    add_table(doc, ["Adım", "Uygulama", "Amaç"], [
        ["1", "STFT: n_fft=2048, hop=512", "Sesin zaman-frekans gösterimini üretmek"],
        ["2", "Yerel spektral tepe bulma", "Baskın ve tekrar edilebilir noktaları seçmek"],
        ["3", "15 × 15 komşuluk; üst %75 tepe eşiği", "Gürültü tabanını azaltmak"],
        ["4", "Saniyede yaklaşık 12 tepe", "Hash sayısını kontrol altında tutmak"],
        ["5", "En fazla 5 hedef; 1–80 frame farkı", "Tepe çiftlerini zaman farkıyla bağlamak"],
        ["6", "SHA-1 tabanlı hash anahtarı", "Kompakt indeks anahtarı üretmek"],
        ["7", "En az 8 hizalanmış hash ve %5 confidence", "Zayıf tesadüfi eşleşmeleri elemek"],
    ], [0.65, 3.35, 2.4])
    add_title(doc, "5.3. Katalog ölçeği ve test mantığı", 2)
    add_body(doc, "Çalışmanın uçak referans tarafında 3.000 civarında referans klipten 851.437 hash üretilmiştir. Kategori fingerprint katmanı genişletildiğinde birleşik katalog yaklaşık 7.364 referans klibe ve 2.585.412 hash'e, kategori tarafı ise 1.733.975 hash'e ulaşmıştır. Bu rakamlar indeks büyüklüğünü anlatır; tek başına doğruluk anlamına gelmez.")
    add_table(doc, ["Fingerprint sonucu", "Kayıt / oran", "Yorum"], [
        ["Uçak referans kataloğu", "3.000 klip / 851.437 hash", "İlk uçak odaklı katalog"],
        ["Birleşik katalog", "7.364 klip / 2.585.412 hash", "Kategori genişlemesi sonrası"],
        ["Kategori hash", "1.733.975", "Ana kategori tarafındaki indeks"],
        ["Uçtan uca kabul", "62 kayıt", "Shazam kabul edilen kayıtların tamamı doğru"],
        ["Fallback", "586 kayıt", "Eşleşme yoksa BEATs/öğrenilmiş yola dönüldü"],
    ], [2.1, 1.65, 2.65])
    add_figure(doc, figures["shazam"], "Şekil 5.1. Parmak izi kataloglarının yaklaşık ölçeği (bin hash; ölçek görselleştirme amaçlıdır).")
    add_title(doc, "5.4. İnsan onayı ve katalog güvenliği", 2)
    add_body(doc, "Agent laboratuvarında model aday etiketi ile altın referans ayrı ayrı dinlenir ve spektrogramları karşılaştırılır. İnsan Onayla, Reddet veya Emin Değilim seçeneklerinden birini seçer. Yalnızca onaylanan kayıtların izole fingerprint indeksine alınması, model önerisinin doğrudan “gerçek” kabul edilmesini engeller. Bu tasarım özellikle yeni fiziksel uçaklar için gereklidir.")
    add_page_break(doc)

    # 6
    add_title(doc, "6. MODEL DENEYLERİ VE KARŞILAŞTIRMALAR", 1)
    add_title(doc, "6.1. Altı üst kategori deneyi", 2)
    add_body(doc, "Ana sınıflandırma deneylerinde modelin yalnızca toplam doğruluğuna bakmak yerine Macro-F1 de izlenmiştir. Sınıf dağılımı dengeli değilse doğruluk tek başına büyük sınıfları ödüllendirebilir. En iyi kaydedilmiş deney, hibrit temsil, dengesizlik sampler'ı ve OTHER uzmanı için focal loss kullanan varyanttır.")
    add_table(doc, ["Yöntem", "Test doğruluğu", "Test Macro-F1", "Karar"], [
        ["Yalnız masking", "%71,60", "%68,34", "Baseline"],
        ["Yalnız contrastive", "%72,10", "%68,82", "Baseline'dan küçük artış"],
        ["Masking + contrastive", "%72,27", "%69,35", "Temsil öğrenimi birleşimi"],
        ["Hibrit + dengesizlik sampler'ı", "%73,28", "%70,66", "Sınıf dengesizliği iyileşti"],
        ["Hibrit + sampler + OTHER uzmanı", "%73,78", "%71,00", "Kaydedilen en iyi varyant"],
    ], [2.8, 1.15, 1.15, 1.35])
    add_figure(doc, figures["main_models"], "Şekil 6.1. Altı üst kategori bağımsız testinde doğruluk ve Macro-F1 karşılaştırması.")
    add_body(doc, "Son uzman deneyinde sınıf bazlı yaklaşık F1 değerleri AIRCRAFT %90, AMBIENT %74, OTHER %59, SPEECH %83, TRAFFIC %78 ve WIND %42 olarak kaydedilmiştir. WIND sınıfında bağımsız test desteği yalnızca 10 kayıt olduğu için bu değer daha geniş bir veriyle tekrar ölçülmelidir.")
    add_title(doc, "6.2. Üç modelin karşılaştırması", 2)
    add_table(doc, ["Model", "Accuracy", "Macro-F1", "Yorum"], [
        ["SVM", "%96,40", "%90,16", "417 kayıtlı karşılaştırmada güçlü ve hafif baseline"],
        ["AirportCNN", "%94,48", "%86,49", "Mel temsiliyle iyi; dış testte genelleme zayıf"],
        ["EfficientNet", "%92,57", "%87,03", "İç karşılaştırmada daha düşük accuracy, dış testte daha dayanıklı"],
    ], [1.8, 1.1, 1.1, 2.45])
    add_title(doc, "6.3. Dış ikili test", 2)
    add_body(doc, "Dış testte iç veri kümesinden farklı dağılım olduğu için model davranışı değişmiştir. AirportCNN %56,48 accuracy ve %36,09 Macro-F1 ile zayıf kalmış, uçak sınıfında precision ve recall sıfır olarak kaydedilmiştir. EfficientNet ise %95,34 accuracy, %94,28 Macro-F1, uçak precision %90,18 ve recall %93,52 üretmiştir. Bu sonuç, tek bir kontrollü testte CNN'in iyi görünmesinin gerçek kullanımda yeterli olmayabileceğini göstermiştir.")
    add_table(doc, ["Dış test", "Accuracy", "Macro-F1", "AIRCRAFT precision", "AIRCRAFT recall"], [
        ["AirportCNN", "%56,48", "%36,09", "%0", "%0"],
        ["EfficientNet", "%95,34", "%94,28", "%90,18", "%93,52"],
    ], [1.75, 1.0, 1.0, 1.35, 1.35])
    add_page_break(doc)

    # 7
    add_title(doc, "7. BAŞARISIZ DENEMELER VE ÇIKARILAN DERSLER", 1)
    add_body(doc, "Bu bölüm özellikle eklenmiştir çünkü proje yalnızca başarılı bir model seçme sürecinden oluşmadı. Bazı sonuçlar validation tarafında iyi görünürken fiziksel olarak bağımsız testte düştü. Bazı teknik fikirler ise doğrudan uygulamada sayısal kararsızlık veya yanlış güven üretti. Bunları saklamak, projeyi olduğundan daha iyi gösterirdi.")
    add_table(doc, ["Deneme", "İlk gözlem", "Bağımsız/uygulama sonucu", "Karar"], [
        ["41 sınıf uçak", "Validation yüksek göründü", "Fiziksel uçak holdout Macro-F1 %13,92", "Üretime alınmadı"],
        ["18 sınıfa daraltma", "Daha savunulabilir split", "Masking %35,24 F1; contrastive %27,99; hybrid %30,77", "Araştırma sonucu"],
        ["Contrastive ilk koşu", "Temsil ayrıştırması hedeflendi", "Numerical overflow / kararsızlık", "Kayıp ölçeği ve normalizasyon gözden geçirildi"],
        ["Çoklu model consensus", "Birden fazla modelin uzlaşması bekleniyordu", "27 kayıt pilotunda consensus %20,83 Macro-F1", "Üretim kararına terfi ettirilmedi"],
        ["CLAP ekli pilot", "Ek genel ses kanıtı", "%25 Macro-F1 ve %59,26 coverage", "Kanıt olarak tutuldu"],
        ["Kapalı küme uçak tahmini", "Model her kayda tip söyleyebildi", "Yeni fiziksel uçakta aşırı güven riski", "Fingerprint/insan kapısı eklendi"],
    ], [1.55, 1.55, 2.45, 1.0], font_size=8.2)
    add_title(doc, "7.1. 41 sınıf uçak denemesinin neden başarısız sayıldığı", 2)
    add_body(doc, "41 sınıflı denemede validation sonucu tek başına yeterli görünse de fiziksel uçak bazlı holdout Macro-F1 yalnızca %13,92 oldu. Bu fark, modelin gerçek uçak kimliğini öğrenmek yerine kayıt koşulunu, kaynak cihazı veya aynı fiziksel örneğin izlerini öğrenmiş olabileceğini gösterdi. Daha sonra en az üç bağımsız fiziksel uçağa sahip sınıflarla 18 sınıfa düşüldü; sonuçlar iyileşti ancak Macro-F1 %35,24 seviyesinde kaldı. Bu yüzden uçak tipi modeli raporda “öneri” olarak sunulmaktadır.")
    add_title(doc, "7.2. Consensus pilotunun neden üretime alınmadığı", 2)
    add_body(doc, "27 kayıtlık consensus pilotunda EfficientNet tek başına %70,37 accuracy ve %70,62 Macro-F1 ile en iyi kanalı oluştururken CNN %20,83, SVM %31,75 ve BEATs 0/0 kapsamla kaldı. En az üç modelin katıldığı consensus sonucu %20,83 Macro-F1 oldu. Burada sorun, model sayısını artırmanın otomatik olarak daha doğru karar üretmemesidir; zayıf veya kapsamı farklı modeller aynı oylamaya eklendiğinde ortak karar daha kötü olabilir. Pilot sonuç bu nedenle üretim yoluna eklenmemiştir.")
    add_title(doc, "7.3. Bu başarısızlıklardan sonra alınan tasarım kararları", 2)
    add_bullets(doc, [
        "Uçak tipinde kapalı küme tahminini kesin kayıt eşleşmesi gibi göstermemek.",
        "Fingerprint kabulünü hash hizalanması, skor ve güven eşikleriyle sınırlamak.",
        "Çoklu model uzlaşmasını ana karar motoru yapmadan önce coverage ve sınıf bazlı hatayı birlikte ölçmek.",
        "Validation başarısını fiziksel kaynak/record bağımsız test olmadan üretim başarısı saymamak.",
        "Model dosyası yoksa arayüzde kanalı sessizce uydurmak yerine kullanılabilir kanallarla kontrollü devam etmek.",
    ])
    add_page_break(doc)

    # 8
    add_title(doc, "8. ÇOKLU PENCERE OYLAMASININ UYGULANMASI VE TEST EDİLMESİ", 1)
    add_title(doc, "8.1. Değişiklikten önceki sorun", 2)
    add_body(doc, "Önceki akışta uzun kayıtlar sabit uzunluklu tek bir parçaya indirgenebildiği için ses olayının kaydın ilerleyen bir bölümünde bulunması karar dışında kalabiliyordu. Bu durum özellikle uçak geçişi gibi kısa ama kaydın belirli bir anında gerçekleşen olaylarda risklidir. Çoklu pencere özelliği bu nedenle yalnızca yeni bir istatistik eklemek değil, veri akışının karar noktasını değiştiren bir uygulama olmuştur.")
    add_title(doc, "8.2. Uygulanan seçim algoritması", 2)
    add_table(doc, ["Parametre", "Değer", "Gerekçe"], [
        ["Pencere uzunluğu", "5 saniye", "Model ve fingerprint ön işleme hattıyla uyum"],
        ["Adım", "2,5 saniye", "Yarı örtüşmeli tarama"],
        ["Maksimum pencere", "5", "Uzun kayıtta hesap maliyetini sınırlamak"],
        ["Kısa kayıt", "Sıfır padding + 1 pencere", "Kayıt kaybetmemek"],
        ["Uzun kayıt", "Temsilci başlangıçlar; son tail zorunlu", "Baş, orta ve son bölümü korumak"],
        ["Kazanan", "Çoğunluk oyu", "Tek pencere aşırı güvenini azaltmak"],
        ["Oy eşitliği", "Ortalama olasılık", "Eşit oyda yumuşak bağlayıcı"],
    ], [1.55, 1.55, 3.35])
    add_figure(doc, figures["window"], "Şekil 8.1. 30,55 saniyelik örnekte temsil edilen beş pencere; her sütun seçilen bir tam pencereyi gösterir.")
    add_body(doc, "Oylama fonksiyonu her pencerenin olasılık satırını önce normalize eder. Her pencerenin en yüksek olasılıklı sınıfı bir oy sayılır. En fazla oyu alan sınıf kazanan olur; iki sınıfın oy sayısı eşitse yalnızca o sınıfların ortalama olasılıkları karşılaştırılır. Böylece tek bir pencerenin çok yüksek confidence değeri, açık çoğunluğu yanlışlıkla tersine çeviremez.")
    add_title(doc, "8.3. Uygulama ayrıntıları", 2)
    add_table(doc, ["Değişen alan", "Uygulama sonucu"], [
        ["window_voting.py", "select_audio_windows ve aggregate_window_probabilities eklendi"],
        ["noise_detector.py", "Ana kategori ML yolu en fazla beş pencereyi çalıştırıyor"],
        ["Metadata", "winner, vote_counts, vote_share, mean_probabilities ve başlangıçlar saklanıyor"],
        ["GUI", "Sonuçta kullanılan pencere sayısı ve oylama bilgisi erişilebilir"],
        ["BEATs hattı", "Seçilen pencerelerde ayrı embedding/head tahmini yapıp aynı oylamayı kullanıyor"],
        ["Tail kontrolü", "Uzun kaydın son tam penceresi seçimde korunuyor"],
    ], [2.0, 4.45])
    add_title(doc, "8.4. Test sonuçları", 2)
    add_table(doc, ["Test grubu", "Sonuç", "Kontrol edilen davranış"], [
        ["Hedefli çoklu pencere testleri", "13/13 başarılı", "Kısa, uzun, tail, padding, çoğunluk, eşitlik ve model entegrasyonu"],
        ["Tam proje testleri", "88/88 başarılı", "Mevcut regresyon testleriyle birlikte"],
        ["Derleme kontrolü", "Başarılı", "noise_detector.py, window_voting.py ve rapor oluşturucu"],
        ["Gerçek uzun WAV", "5 pencere", "30,55 s traffic.wav; başlangıçlar 0, 7,5, 15, 20, 25,55 s"],
    ], [1.9, 1.2, 3.35])
    add_figure(doc, figures["tests"], "Şekil 8.2. Çoklu pencere değişikliğinin hedefli ve tam test özetleri.")
    add_note(doc, "Gerçek traffic.wav denemesi mekanizmanın beş pencere ürettiğini gösterir; tek başına genel doğruluk testi değildir. Bu ayrım önemlidir: oylama algoritmasının çalışması ile oylamanın her ses sınıfında doğruluğu artırması farklı iddialardır.")
    add_page_break(doc)

    # 9
    add_title(doc, "9. UÇTAN UCA UYGULAMA DAVRANIŞI VE KULLANICI ARAYÜZÜ", 1)
    add_title(doc, "9.1. İki aşamalı kullanıcı akışı", 2)
    add_body(doc, "Ana arayüzde kullanıcı önce “Üst Sınıfı Bul” adımıyla sesin ana kategorisini bulur. Daha sonra “Shazam / Alt Türü Bul” adımıyla AIRCRAFT için parmak izi eşleşmesi, TRAFFIC ve OTHER için alt tür modeli çalıştırılır. Bu iki butonun ayrılması, modelin henüz üst sınıfı bilmeden uçak tipi söylemesini önler.")
    add_table(doc, ["Kullanıcı adımı", "Sistem davranışı", "Kullanıcıya gösterilen kanıt"], [
        ["Ses seç / kaydet", "Dosya veya mikrofon girdisi alınır", "Dosya adı, süre ve analiz durumu"],
        ["Üst Sınıfı Bul", "Pencereler ayrı işlenir ve oylanır", "Ana sınıf, confidence, oy sayıları"],
        ["Uçak kontrolü", "AIRCRAFT değilse alt uçak yolu durur", "Guard sonucu / açıklama"],
        ["Shazam / Alt Türü Bul", "Fingerprint veya alt tür modeli seçilir", "Yöntem etiketi ve alt tür"],
        ["İnsan onayı", "Agent laboratuvarında karar saklanır", "Onay, red veya emin değilim"],
    ], [1.45, 2.95, 2.05])
    add_title(doc, "9.2. Uçtan uca 648 kayıtlık sonuç", 2)
    add_body(doc, "648 kayıtlık uçtan uca ölçümde ana kategori 621 kayıtta doğru olup %95,83 sonucuna ulaşmıştır. Ana kategori yanlışsa alt tür de doğal olarak etkilenir. Shazam kabul edilen 62 kaydın tamamı doğru olmuş, 586 kayıt fallback yoluna gitmiştir. Alt tür/entegre sonuç v1'de 537/648 (%82,87), v2'de 561/648 (%86,57) olarak kaydedilmiştir. AIRCRAFT false positive değerinin sıfır olması, guard katmanının uçak olmayan kayıtları uçak olarak etiketlememe hedefinde yararlı olduğunu göstermektedir.")
    add_table(doc, ["Uçtan uca ölçüm", "Doğru / toplam", "Oran", "Açıklama"], [
        ["Ana kategori", "621 / 648", "%95,83", "Üst sınıf sonucu"],
        ["Shazam kabul", "62 / 62", "%100", "Kabul edilen fingerprint sonuçları"],
        ["Alt tür / v1", "537 / 648", "%82,87", "Eski alt akış"],
        ["Alt tür / v2", "561 / 648", "%86,57", "Güncellenmiş alt akış"],
        ["AIRCRAFT false positive", "0", "0", "Guard sonrası yanlış uçak kabulü"],
        ["Fallback", "586 / 648", "%90,43", "Shazam yerine öğrenilmiş yol"],
    ], [2.0, 1.25, 0.9, 2.3])
    add_title(doc, "9.3. İnsan denetimi", 2)
    add_body(doc, "Bu proje tamamen otomatik karar veren bir sistem olarak sunulmamıştır. Özellikle fingerprint kataloğuna yeni kayıt eklemek için insanın sesleri ayrı ayrı dinlemesi ve referansla karşılaştırması gerekir. Arayüzde tahmin, kullanılan yöntem ve güveni ayrı göstermek; kullanıcıya model önerisini kesin gerçek gibi sunmamak için bilinçli bir tercihtir.")
    add_note(doc, "Uçtan uca sonuçlar uygulamanın bütün karar zincirini ölçer; yalnızca modelin tek başına test metriği değildir. Bu nedenle Bölüm 6'daki Macro-F1 tabloları ile bu bölümdeki 648 kayıtlık oranlar birbirinin yerine kullanılmamalıdır.")
    add_page_break(doc)

    # 10
    add_title(doc, "10. YENİDEN ÜRETİLEBİLİRLİK, SINIRLILIKLAR VE RİSKLER", 1)
    add_title(doc, "10.1. Kurulum ve çalıştırma", 2)
    add_body(doc, "Desteklenen temel ortam Windows 10/11 ve Python 3.10/3.11'dir. GUI ve hafif model kanalları CPU üzerinde çalışabilir; EfficientNet, BEATs ve eğitim adımları için NVIDIA GPU önerilir. Temel kurulumda sanal ortam oluşturulup requirements.txt kurulmalıdır. Model ağırlıkları, ham sesler, embedding önbellekleri ve SQLite indeksleri boyut/lisans sebebiyle Git'e eklenmediğinden, temiz klondan sonra bunların ayrıca hazırlanması gerekir.")
    add_table(doc, ["İşlem", "Komut / dosya", "Beklenen çıktı"], [
        ["Temel kurulum", "python -m venv venv; pip install -r requirements.txt", "Çalışma ortamı"],
        ["Ana arayüz", "python app_shazam.py", "Kategori + Shazam/alt tür GUI"],
        ["Agent", "python app_agent.py", "Çoklu model ve insan-onay laboratuvarı"],
        ["Veritabanı", "python app_database.py", "Salt-okunur SQLite kanıt ekranı"],
        ["Pencere testleri", "python -m unittest -v test_window_voting", "Pencereleme ve oylama testleri"],
        ["Derleme", "python -m compileall -q ...", "Sözdizimi doğrulaması"],
    ], [1.45, 3.35, 1.65])
    add_title(doc, "10.2. Bilinen sınırlılıklar", 2)
    add_table(doc, ["Sınırlılık", "Etkisi", "Azaltma önerisi"], [
        ["Veri miktarı ve kaynak çeşitliliği", "Yeni fiziksel uçakta genelleme düşer", "Daha fazla bağımsız kayıt ve grup bazlı split"],
        ["WIND test desteği 10 kayıt", "Sınıf F1'i oynak", "Daha geniş bağımsız test"],
        ["Uçak tipi Macro-F1 %50,20", "Kesin tip iddiası yapılamaz", "Daha çok fiziksel uçak + open-set eşik"],
        ["Yerel artefaktlar eksik olabilir", "Temiz klon tüm sonucu çalıştıramaz", "Sürümlü model/indeks deposu veya LFS"],
        ["Beş pencere hesap maliyeti", "Tek pencereye göre süre artar", "Model batch inference ve adaptive sampling"],
        ["Kapalı kategori varsayımı", "Tanımsız ses zorla bilinen sınıfa gidebilir", "Reject/unknown sınıfı ve kalibrasyon"],
        ["Fingerprint katalog kapsamı", "Katalog dışı uçak kesin eşleşemez", "İnsan onaylı referans laboratuvarı"],
    ], [1.65, 2.4, 2.4])
    add_title(doc, "10.3. Güvenlik ve yanlış güven riski", 2)
    add_body(doc, "Bir modelin confidence değerinin yüksek olması onun gerçek hayatta doğru olduğunu garanti etmez. Özellikle kapalı küme uçak tipi modelinde yeni bir ses, mevcut sekiz sınıftan birine yüksek olasılıkla zorlanabilir. Projede bu riski azaltmak için fingerprint ve insan onayı ayrılmış, ana kategori guard'ı eklenmiş ve sonuçta yöntem bilgisi saklanmıştır. Gelecek sürümde reject/unknown sınıfı ve olasılık kalibrasyonu üretim güvenini daha da artıracaktır.")
    add_page_break(doc)

    # 11
    add_title(doc, "11. SONUÇ VE GELECEK ÇALIŞMA PLANI", 1)
    add_title(doc, "11.1. Proje sonunda ne elde edildi?", 2)
    add_body(doc, "Proje sonunda yalnızca bir sınıflandırma modeli değil, farklı kanıt türlerini birleştiren çalışabilir bir prototip ortaya çıktı. Ana kategori tarafında en iyi bağımsız test sonucu %73,78 accuracy ve %71,00 Macro-F1 olarak kaydedildi; uygulama zincirinde 648 kayıtlık ana kategori sonucu %95,83 oldu. TRAFFIC alt türü %98,14 Macro-F1 ile güçlü, OTHER alt türü %79,04 ile kullanılabilir fakat geliştirmeye açık, uçak tipi modeli ise %50,20 Macro-F1 ile açıkça sınırlı kaldı. Bu tablo, sistemin her alt problemde aynı olgunlukta olmadığını gösteriyor.")
    add_body(doc, "En önemli mühendislik çıktılarından biri çoklu pencere oylamasıdır. Bu değişiklikle uzun kayıtta yalnızca tek bir bölümün karar vermesi engellendi; seçilen pencereler, pencere başlangıçları, oy sayıları ve ortalama olasılıklar sonuç metadatasına taşındı. 13 hedefli ve 88 tam testin başarılı olması değişikliğin mevcut akışta regresyon oluşturmadığını gösterdi. Bununla birlikte beş pencerenin her veri kümesinde doğruluğu artırdığı iddiası henüz kanıtlanmış değildir; bu konu ayrı bir kontrollü ablation deneyi gerektirir.")
    add_title(doc, "11.2. Önceliklendirilmiş gelecek işler", 2)
    add_table(doc, ["Öncelik", "İş", "Neden gerekli?", "Beklenen katkı"], [
        ["1", "Open-set / unknown uçak kararı", "Yeni uçağı zorla bilinen tipe atamamak", "Yanlış güven azalır"],
        ["2", "Fiziksel uçak sayısını artırmak", "41/18 sınıf denemelerinde genelleme düşük", "Uçak tipi F1 yükselir"],
        ["3", "Çoklu pencere ablation", "1, 3 ve 5 pencereyi aynı veriyle ölçmek", "Oylamanın gerçek katkısı ayrışır"],
        ["4", "Büyük ve dengeli WIND/OTHER veri", "Zayıf sınıf desteğini artırmak", "Macro-F1 daha güvenilir olur"],
        ["5", "Model kalibrasyonu", "Confidence ile gerçek hata oranını yaklaştırmak", "GUI güveni daha anlamlı olur"],
        ["6", "Model/indeks paketleme", "Temiz kurulumun sonuç üretebilmesi", "Teslim ve yeniden üretim kolaylaşır"],
    ], [0.65, 2.25, 2.3, 1.25])
    add_title(doc, "11.3. Son değerlendirme", 2)
    add_body(doc, "Bu çalışmada başarılı olan taraf, farklı ses işleme fikirlerini tek bir arayüzde birleştirirken kanıt türlerini birbirine karıştırmamaktır. Başarısız olan taraf ise özellikle yeni fiziksel uçak tipini yalnızca sınıflandırıcıyla güvenilir biçimde genelleme hedefidir. Bu başarısızlık projenin değersiz olduğunu değil, gerçek kullanım için hangi verinin ve hangi değerlendirme düzeninin eksik olduğunu açıkça gösterir. Raporun sonraki okuyucusu sistemi çalıştırmak istediğinde bu ayrımı bilirse, sonuçları yanlış yorumlamadan projeyi geliştirebilir.")
    add_page_break(doc)

    # 12
    add_title(doc, "12. TERİMLER, DOSYALAR VE KAYNAK NOTLARI", 1)
    add_title(doc, "12.1. Kısa terim sözlüğü", 2)
    add_table(doc, ["Terim", "Bu projedeki anlamı"], [
        ["Macro-F1", "Her sınıfın F1 değerinin aritmetik ortalaması; sınıf dengesizliğine duyarlı özet"],
        ["Accuracy", "Doğru tahminlerin tüm örneklere oranı"],
        ["Fingerprint", "Spektral tepe çiftlerinden oluşturulan ve kayıt arayan akustik imza"],
        ["Hash hizalanması", "Sorgu ve referans hash'lerinin zaman farklarının tutarlı olması"],
        ["Fallback", "Öncelikli fingerprint eşleşmesi yoksa öğrenilmiş alt tür kanalına dönüş"],
        ["Guard", "Bir sonraki karar yolunu yalnızca uygun ana sınıf için açan güvenlik kontrolü"],
        ["Physical-aircraft split", "Aynı fiziksel uçağın splitler arasında sızıntı yapmaması"],
        ["Open-set", "Eğitimde olmayan sınıfı unknown olarak reddedebilme yaklaşımı"],
        ["Vote share", "Kazanan sınıfın tüm pencerelerde aldığı oy oranı"],
    ], [1.8, 4.65])
    add_title(doc, "12.2. Başvurulan proje içi kaynaklar", 2)
    add_table(doc, ["Kaynak", "Rapor içinde kullanıldığı yer"], [
        ["README.md", "Mimari, kurulum, model/artefakt düzeni, ana metrikler"],
        ["noise_detector.py", "Ana kategori ve çoklu model orkestrasyonu"],
        ["window_voting.py", "Pencere seçimi ve oy birleştirme algoritması"],
        ["test_window_voting.py", "Çoklu pencere hedefli testleri ve entegrasyon testleri"],
        ["aircraft_fingerprint.py", "STFT, tepe ve hash tabanlı eşleştirme"],
        ["gui_main.py", "Masaüstü kullanıcı akışı"],
        ["COLAB_UCAK_ALT_TUR_EGITIMI.md", "Fiziksel uçak bağımsızlığı ve veri yeterliliği notları"],
        ["Örnek_Rapor_ (3).docx", "Yalnızca rapor biçimi ve ayrıntı seviyesi için örnek"],
    ], [2.55, 3.9])
    add_title(doc, "12.3. Son söz", 2)
    add_body(doc, "Bu belgeyi okuyan kişinin projeyi anlaması için özellikle üç sınırın akılda tutulması gerekir: fingerprint yalnızca katalogda doğrulanmış kayıtları arar; uçak tipi öğrenilmiş modeli yeni fiziksel uçaklarda kesin değildir; çoklu pencere oylaması uygulandı ve test edildi ancak doğruluk katkısı ayrıca ölçülmelidir. Bu üç sınır, sistemin bugün ne yaptığını ve sonraki çalışmanın nereden başlaması gerektiğini açıklar.")
    add_note(doc, "Rapor tarihi ve kapak bilgileri teslimden önce öğrenci adı, bölüm, kurum ve staj tarihleriyle doldurulmalıdır.")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.core_properties.title = "Havalimanı Gürültü Tespit ve Ses Tanıma Sistemi - Staj Projesi Raporu"
    doc.core_properties.subject = "Çoklu pencere oylaması, akustik parmak izi ve ses sınıflandırma"
    doc.core_properties.author = "[Ad Soyad]"
    doc.core_properties.keywords = "havalimanı, gürültü, ses sınıflandırma, Shazam, çoklu pencere, staj"
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    result = build_report()
    print(result)
